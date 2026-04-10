import streamlit as st
import docx
import os
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import io

# ==========================================
# CẤU HÌNH TRANG WEB
# ==========================================
st.set_page_config(page_title="Tạo Đề Cương Luận Văn", page_icon="🎓", layout="centered")

st.title("🎓 Trình Tạo Đề Cương Luận Văn Chuẩn")
st.write("Hệ thống tự động dàn trang, thuật toán CĂN ĐỀU mốc trên/dưới, chèn LOGO, tạo MỤC LỤC và số trang kép.")
st.divider()

# ==========================================
# CÁC HÀM CAN THIỆP XML & THUẬT TOÁN DÀN TRANG
# ==========================================
def calculate_even_spaces(total_allowed_lines, used_lines, num_gaps):
    """Thuật toán chia đều khoảng trống: Lấy tổng không gian trừ đi không gian chữ, sau đó chia đều cho các khe hở"""
    remaining = max(num_gaps, total_allowed_lines - used_lines) # Đảm bảo mỗi khe luôn >= 1
    base = remaining // num_gaps
    remainder = remaining % num_gaps
    
    spaces = [base] * num_gaps
    # Nếu chia không hết, cộng dồn phần dư xuống các khe phía đáy trang
    for i in range(remainder):
        spaces[-(i+1)] += 1
    return spaces

def add_page_border(sect_pr):
    borders = OxmlElement('w:pgBorders')
    borders.set(qn('w:offsetFrom'), 'text')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12') 
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    sect_pr.append(borders)

def clear_page_border(sect_pr):
    for borders in sect_pr.xpath('./w:pgBorders'):
        sect_pr.remove(borders)

def add_page_number(paragraph):
    p = paragraph._p
    run = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run.append(fldChar1)
    run.append(instrText)
    run.append(fldChar2)
    run.append(fldChar3)
    p.append(run)

def setup_toc_styles(doc):
    for i in range(1, 4):
        style_name = f'TOC {i}'
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)
        style.font.bold = False 
        if i == 1:
            style.paragraph_format.left_indent = Cm(0)
        elif i == 2:
            style.paragraph_format.left_indent = Cm(1.27) 
        elif i == 3:
            style.paragraph_format.left_indent = Cm(2.54)

def add_toc_to_doc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MỤC LỤC")
    r.bold, r.font.name, r.font.size = True, 'Times New Roman', Pt(14)
    
    p_trang = doc.add_paragraph()
    p_trang.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_trang = p_trang.add_run("Trang")
    r_trang.bold, r_trang.font.name, r_trang.font.size = True, 'Times New Roman', Pt(13)
    
    p_toc = doc.add_paragraph()
    run = p_toc.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_note = p_note.add_run("\n[HƯỚNG DẪN HIỂN THỊ MỤC LỤC VÀ CẬP NHẬT SỐ TRANG]\n1. Nhấn nút 'Enable Editing' (Bật chỉnh sửa) màu vàng ở phía trên cùng màn hình Word.\n2. Bấm tổ hợp phím Ctrl + P (để Word nhận diện số trang), sau đó bấm Esc để quay lại.\n3. Nhấp CHUỘT PHẢI vào dòng chữ đỏ này -> Chọn 'Update Field' -> Chọn 'Update entire table' -> OK.\n")
    r_note.font.name, r_note.font.size, r_note.font.italic = 'Times New Roman', Pt(11), True
    r_note.font.color.rgb = RGBColor(255, 0, 0)

def set_pgnum_type(sectPr, fmt='decimal', start='1'):
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), start)
    
    cols = sectPr.xpath('./w:cols')
    docGrid = sectPr.xpath('./w:docGrid')
    if cols:
        cols[0].addprevious(pgNumType)
    elif docGrid:
        docGrid[0].addprevious(pgNumType)
    else:
        sectPr.append(pgNumType)

# ==========================================
# HÀM TẠO BẢNG DANH MỤC
# ==========================================
def create_two_col_table(doc, col1_name, col2_name):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = col1_name
    hdr_cells[1].text = col2_name
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.0)
    for _ in range(3):
        table.add_row()

# ==========================================
# HÀM ĐỆ QUY TẠO GIAO DIỆN NHẬP LIỆU
# ==========================================
def render_section(level, prefix, key_prefix):
    with st.container(border=True):
        st.markdown(f"**Mục {prefix}**")
        title = st.text_input("Tên mục:", key=f"title_{key_prefix}", label_visibility="collapsed", placeholder=f"Tên mục {prefix}")
        content = st.text_area("Nội dung:", key=f"content_{key_prefix}", height=100, label_visibility="collapsed", placeholder=f"Nội dung mục {prefix}")
        children = []
        if level < 4:
            num_children = st.number_input(f"Số tiểu mục con trong {prefix}:", min_value=0, max_value=15, value=0, step=1, key=f"num_{key_prefix}")
            for k in range(int(num_children)):
                child_prefix = f"{prefix}.{k+1}"
                children.append(render_section(level+1, child_prefix, f"{key_prefix}_{k}"))
    return {"title": title, "content": content, "children": children}

def apply_academic_rules(node):
    if node.get("children"):
        pruned_children = [apply_academic_rules(c) for c in node["children"]]
        if len(pruned_children) == 1:
            single_child = pruned_children[0]
            merged_text = single_child["title"]
            if single_child["content"].strip():
                merged_text += "\n" + single_child["content"]
            if node.get("content", "").strip():
                node["content"] += "\n\n" + merged_text
            else:
                node["content"] = merged_text
            node["children"] = single_child["children"]
        else:
            node["children"] = pruned_children
    return node

# ==========================================
# GIAO DIỆN NHẬP LIỆU CHÍNH
# ==========================================
st.header("I. THÔNG TIN BÌA")
thesis_title = st.text_input("Tên Đề tài Luận văn:", placeholder="Ví dụ: HIỆU QUẢ CỦA CHƯƠNG TRÌNH CAN THIỆP HABIT-ILE...")
author_name = st.text_input("Họ và tên tác giả:", placeholder="Ví dụ: TRẦN MINH HOÀNG")

st.markdown("##### Người hướng dẫn khoa học")
supervisor_1 = st.text_input("1. Họ và tên người hướng dẫn 1:", placeholder="Ví dụ: PGS.TS. NGUYỄN VĂN A")
supervisor_2 = st.text_input("2. Họ và tên người hướng dẫn 2 (Bỏ trống nếu không có):", placeholder="Ví dụ: TS. TRẦN THỊ B")

st.divider()
st.header("II. NỘI DUNG")

st.subheader("ĐẶT VẤN ĐỀ")
dat_van_de_content = st.text_area("Nội dung phần Đặt vấn đề:", height=200, key="dvd")

# --- CHƯƠNG 1 ---
st.markdown("### Chương 1. TỔNG QUAN TÀI LIỆU")
c1_intro = st.text_area("Nội dung dẫn nhập Chương 1 (nếu có):", height=100, key="c1_intro")
c1_num = st.number_input("Số lượng mục cấp 2 (ví dụ: 1.1, 1.2):", min_value=0, max_value=20, value=2, step=1, key="c1_num")
c1_children = [render_section(2, f"1.{j+1}", f"c1_sec_{j}") for j in range(c1_num)]
st.write("---")

# --- CHƯƠNG 2 (CỐ ĐỊNH 9 MỤC) ---
st.markdown("### Chương 2. PHƯƠNG PHÁP NGHIÊN CỨU")
c2_intro = st.text_area("Nội dung dẫn nhập Chương 2 (nếu có):", height=100, key="c2_intro")
c2_fixed_titles = [
    "Thiết kế nghiên cứu", "Thời gian và địa điểm nghiên cứu", "Đối tượng nghiên cứu", 
    "Cỡ mẫu của nghiên cứu", "Xác định các biến số độc lập và phụ thuộc", 
    "Phương pháp và công cụ đo lường, thu thập số liệu", "Quy trình nghiên cứu", 
    "Phương pháp phân tích dữ liệu", "Đạo đức trong nghiên cứu"
]
c2_children = []
for j, title in enumerate(c2_fixed_titles):
    with st.expander(f"Mục 2.{j+1}. {title}", expanded=True):
        c2_content = st.text_area(f"Nội dung mục 2.{j+1}:", height=150, key=f"c2_sec_{j}")
        c2_children.append({"title": title, "content": c2_content, "children": []})
st.write("---")

# --- CHƯƠNG 3 ---
st.markdown("### Chương 3. DỰ KIẾN KẾT QUẢ")
c3_intro = st.text_area("Nội dung dẫn nhập Chương 3 (nếu có):", height=100, key="c3_intro")
c3_num = st.number_input("Số lượng mục cấp 2 (ví dụ: 3.1, 3.2):", min_value=0, max_value=20, value=2, step=1, key="c3_num")
c3_children = [render_section(2, f"3.{j+1}", f"c3_sec_{j}") for j in range(c3_num)]
st.write("---")

# --- CHƯƠNG 4 ---
st.markdown("### Chương 4. KẾ HOẠCH THỰC HIỆN")
c4_intro = st.text_area("Nội dung dẫn nhập Chương 4 (nếu có):", height=100, key="c4_intro")
c4_num = st.number_input("Số lượng mục cấp 2 (ví dụ: 4.1, 4.2):", min_value=0, max_value=20, value=1, step=1, key="c4_num")
c4_children = [render_section(2, f"4.{j+1}", f"c4_sec_{j}") for j in range(c4_num)]
st.write("---")

# Các phần cuối
st.subheader("DANH MỤC CÁC CÔNG TRÌNH CÔNG BỐ CÓ LIÊN QUAN")
danh_muc_content = st.text_area("Nội dung (Nếu không có hãy để trống):", height=150, key="dm")
st.subheader("TÀI LIỆU THAM KHẢO")
tai_lieu_content = st.text_area("Danh sách tài liệu tham khảo:", height=200, key="tl")
st.subheader("PHỤ LỤC")
phu_luc_content = st.text_area("Nội dung Phụ lục (Nếu có):", height=200, key="pl")
st.divider()

# ==========================================
# CÁC HÀM HỖ TRỢ XUẤT FILE WORD 
# ==========================================
def add_empty_lines(doc, num_lines, size=16):
    """Tạo dòng trống giãn dòng chuẩn 1.5"""
    if num_lines > 0:
        for _ in range(int(num_lines)):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run()
            r.font.size = Pt(size)

def add_cover_para(doc, text, size=16, bold=True):
    """Hàm tạo dòng chữ ở trang bìa giãn dòng chuẩn 1.5 (Không dư khoảng đệm)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    return p

def add_main_heading(doc, text):
    try:
        p = doc.add_paragraph(style='Heading 1')
    except KeyError:
        p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0] if p.runs else p.add_run()
    r.text = text 
    r.bold, r.font.name, r.font.size = True, 'Times New Roman', Pt(14)
    r.font.color.rgb = RGBColor(0, 0, 0) 

def add_normal_text(doc, text_content):
    if not text_content.strip(): return
    for para_text in text_content.split('\n'):
        if para_text.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(1.27)
            r = p.add_run(para_text.strip())
            r.font.name, r.font.size = 'Times New Roman', Pt(13)

def write_sections_to_word(doc, children_list, prefix_list):
    for i, child in enumerate(children_list):
        current_prefix = prefix_list + [str(i + 1)]
        prefix_str = ".".join(current_prefix)
        level = len(current_prefix)
        style_name = f'Heading {level}' if level <= 3 else 'Heading 3'
        
        if child["title"].strip() or child["content"].strip() or child["children"]:
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0) 
            title_text = f"{prefix_str}. {child['title']}" if child['title'].strip() else f"{prefix_str}."
            r = p.runs[0] if p.runs else p.add_run()
            r.text = title_text
            r.bold, r.font.name, r.font.size = True, 'Times New Roman', Pt(13)
            r.font.color.rgb = RGBColor(0, 0, 0)
            
            add_normal_text(doc, child["content"])
            if child.get("children"):
                write_sections_to_word(doc, child["children"], current_prefix)


# ==========================================
# NÚT XỬ LÝ VÀ TẠO FILE WORD
# ==========================================
if st.button("✨ TẠO FILE WORD HOÀN CHỈNH", type="primary", use_container_width=True):
    if not thesis_title or not supervisor_1:
        st.warning("⚠️ Vui lòng nhập Tên Đề tài và ít nhất 1 Người hướng dẫn!")
    else:
        with st.spinner("Đang kích hoạt thuật toán CĂN ĐỀU mốc không gian và dàn trang..."):
            c1_processed = {"title": "TỔNG QUAN TÀI LIỆU", "content": c1_intro, "children": [apply_academic_rules(c) for c in c1_children]}
            c2_processed = {"title": "PHƯƠNG PHÁP NGHIÊN CỨU", "content": c2_intro, "children": c2_children}
            c3_processed = {"title": "DỰ KIẾN KẾT QUẢ", "content": c3_intro, "children": [apply_academic_rules(c) for c in c3_children]}
            c4_processed = {"title": "KẾ HOẠCH THỰC HIỆN", "content": c4_intro, "children": [apply_academic_rules(c) for c in c4_children]}
            all_chapters = [c1_processed, c2_processed, c3_processed, c4_processed]

            doc = docx.Document()
            setup_toc_styles(doc)
            
            style_normal = doc.styles['Normal']
            style_normal.font.name, style_normal.font.size = 'Times New Roman', Pt(13)

            try:
                element_updatefields = OxmlElement('w:updateFields')
                element_updatefields.set(qn('w:val'), 'true')
                doc.settings.element.append(element_updatefields)
            except Exception:
                pass

            title_lines = (len(thesis_title) // 40) + 1
            has_logo = os.path.exists("logo_UMP.png")
            has_sup2 = bool(supervisor_2.strip())

            # =====================================
            # SECTION 1: TRANG BÌA CHÍNH (CÓ LOGO)
            # =====================================
            sec_0 = doc.sections[0]
            sec_0.top_margin, sec_0.bottom_margin = Cm(3.5), Cm(3.0)
            sec_0.left_margin, sec_0.right_margin = Cm(3.5), Cm(2.0)
            add_page_border(sec_0._sectPr)

            # --- TÍNH TOÁN CĂN ĐỀU BÌA 1 ---
            # Số dòng giấy A4 tối đa quy ước an toàn: 23 dòng (đã trừ Margin)
            # Khối chữ cố định đã dùng: BỘ (1) + ĐẠI HỌC (1) + Logo(4) + TÁC GIẢ(1) + TITLE(n) + ĐỀ CƯƠNG(1) + TPHCM(1)
            used_lines_1 = 5 + title_lines + (4 if has_logo else 0)
            # Tổng số khe hở giữa ĐẠI HỌC và TPHCM: 4 khe (Khe_Logo, Khe_TácGiả, Khe_TiêuĐề, Khe_ĐềCương)
            spaces_1 = calculate_even_spaces(23, used_lines_1, 4)

            # Header Bìa 1 (2 dòng sát nhau không cách)
            table = doc.add_table(rows=1, cols=2)
            p_left = table.cell(0, 0).paragraphs[0]
            p_left.paragraph_format.space_after = Pt(0)
            p_left.paragraph_format.line_spacing = 1.5
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r_left = p_left.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO")
            r_left.font.name, r_left.font.size = 'Times New Roman', Pt(16)
            
            p_right = table.cell(0, 1).paragraphs[0]
            p_right.paragraph_format.space_after = Pt(0)
            p_right.paragraph_format.line_spacing = 1.5
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_right = p_right.add_run("BỘ Y TẾ")
            r_right.font.name, r_right.font.size = 'Times New Roman', Pt(16)
            
            # --- MỐC TRÊN ---
            add_cover_para(doc, "ĐẠI HỌC Y DƯỢC THÀNH PHỐ HỒ CHÍ MINH", 16, True)

            # Khe hở 1 (chia đều khoảng Logo)
            if has_logo:
                add_empty_lines(doc, 1, 16)
                try:
                    p_logo = doc.add_paragraph()
                    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_logo.paragraph_format.space_after = Pt(0)
                    p_logo.paragraph_format.line_spacing = 1.5
                    r_logo = p_logo.add_run()
                    r_logo.add_picture("logo_UMP.png", width=Cm(3.5)) 
                except Exception:
                    pass
                add_empty_lines(doc, max(1, spaces_1[0] - 1), 16)
            else:
                add_empty_lines(doc, spaces_1[0], 16)

            add_cover_para(doc, author_name.upper(), 16, True)
            
            # Khe hở 2
            add_empty_lines(doc, spaces_1[1], 16)
            add_cover_para(doc, thesis_title.upper(), 20, True)

            # Khe hở 3
            add_empty_lines(doc, spaces_1[2], 16)
            add_cover_para(doc, "ĐỀ CƯƠNG LUẬN VĂN THẠC SĨ", 16, True)

            # Khe hở 4
            add_empty_lines(doc, spaces_1[3], 16)
            
            # --- MỐC DƯỚI ---
            add_cover_para(doc, "THÀNH PHỐ HỒ CHÍ MINH - NĂM 2026", 16, True)


            # =====================================
            # SECTION 2: TRANG BÌA PHỤ (KHÔNG LOGO)
            # =====================================
            new_section_cover_2 = doc.add_section(WD_SECTION.NEW_PAGE)
            new_section_cover_2.top_margin, new_section_cover_2.bottom_margin = Cm(3.5), Cm(3.0)
            new_section_cover_2.left_margin, new_section_cover_2.right_margin = Cm(3.5), Cm(2.0)
            add_page_border(new_section_cover_2._sectPr)

            # --- TÍNH TOÁN CĂN ĐỀU BÌA 2 ---
            # Khối chữ cố định đã dùng: BỘ(1) + ĐẠI HỌC(1) + TÁC GIẢ(1) + TITLE(n) + NGÀNH(1) + MÃ SỐ(1) + ĐỀ CƯƠNG(1) + NGƯỜI HD(1) + TÊN_HD(1/2) + TPHCM(1)
            used_lines_2 = 9 + title_lines + (1 if has_sup2 else 0)
            # Tổng số khe hở giữa ĐẠI HỌC và TPHCM: 5 khe (Khe_TácGiả, Khe_TiêuĐề, Khe_KhốiNgành, Khe_ĐềCương, Khe_KhốiHD)
            spaces_2 = calculate_even_spaces(25, used_lines_2, 5)

            # Header Bìa 2 (2 dòng sát nhau không cách)
            table = doc.add_table(rows=1, cols=2)
            p_left = table.cell(0, 0).paragraphs[0]
            p_left.paragraph_format.space_after = Pt(0)
            p_left.paragraph_format.line_spacing = 1.5
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r_left = p_left.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO")
            r_left.font.name, r_left.font.size = 'Times New Roman', Pt(16)
            
            p_right = table.cell(0, 1).paragraphs[0]
            p_right.paragraph_format.space_after = Pt(0)
            p_right.paragraph_format.line_spacing = 1.5
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_right = p_right.add_run("BỘ Y TẾ")
            r_right.font.name, r_right.font.size = 'Times New Roman', Pt(16)
            
            # --- MỐC TRÊN ---
            add_cover_para(doc, "ĐẠI HỌC Y DƯỢC THÀNH PHỐ HỒ CHÍ MINH", 16, True)

            # Khe hở 1
            add_empty_lines(doc, spaces_2[0], 16)
            add_cover_para(doc, author_name.upper(), 16, True)
            
            # Khe hở 2
            add_empty_lines(doc, spaces_2[1], 16)
            add_cover_para(doc, thesis_title.upper(), 20, True)

            # Khe hở 3
            add_empty_lines(doc, spaces_2[2], 16)
            # --- KHỐI GỘP: NGÀNH & MÃ SỐ (Không cách dòng) ---
            add_cover_para(doc, "NGÀNH: KỸ THUẬT PHỤC HỒI CHỨC NĂNG", 16, True)
            add_cover_para(doc, "MÃ SỐ: 8720603", 16, True)

            # Khe hở 4
            add_empty_lines(doc, spaces_2[3], 16)
            add_cover_para(doc, "ĐỀ CƯƠNG LUẬN VĂN THẠC SĨ", 16, True)

            # Khe hở 5
            add_empty_lines(doc, spaces_2[4], 16)
            # --- KHỐI GỘP: NGƯỜI HD & TÊN (Không cách dòng) ---
            add_cover_para(doc, "NGƯỜI DỰ KIẾN HƯỚNG DẪN KHOA HỌC:", 16, True)
            if not has_sup2:
                add_cover_para(doc, f"{supervisor_1.upper()}", 16, True)
            else:
                add_cover_para(doc, f"1. {supervisor_1.upper()}", 16, True)
                add_cover_para(doc, f"2. {supervisor_2.upper()}", 16, True)

            # Khe phân bổ phụ dư của hàm chia đều (để đẩy đáy xuống sát)
            add_empty_lines(doc, 1 + spaces_2[4] // 2, 16)
            
            # --- MỐC DƯỚI ---
            add_cover_para(doc, "THÀNH PHỐ HỒ CHÍ MINH - NĂM 2026", 16, True)


            # =====================================
            # SECTION 3: CÁC TRANG DANH MỤC ĐỆM (SỐ LA MÃ)
            # =====================================
            new_section_prelim = doc.add_section(WD_SECTION.NEW_PAGE)
            clear_page_border(new_section_prelim._sectPr)
            
            new_section_prelim.header.is_linked_to_previous = False
            header_para = new_section_prelim.header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_page_number(header_para) 
            set_pgnum_type(new_section_prelim._sectPr, fmt='lowerRoman', start='1')
            
            add_toc_to_doc(doc)
            doc.add_page_break()

            add_main_heading(doc, "DANH MỤC CÁC TỪ VIẾT TẮT")
            create_two_col_table(doc, "Từ viết tắt", "Ý nghĩa")
            doc.add_page_break()

            add_main_heading(doc, "DANH MỤC ĐỐI CHIẾU CÁC THUẬT NGỮ ANH - VIỆT")
            create_two_col_table(doc, "Tiếng Anh", "Tiếng Việt")
            doc.add_page_break()

            add_main_heading(doc, "DANH MỤC CÁC BẢNG")
            p_b = doc.add_paragraph("(Chèn danh mục bảng tự động tại đây bằng Word)")
            p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

            add_main_heading(doc, "DANH MỤC CÁC HÌNH")
            p_h = doc.add_paragraph("(Chèn danh mục hình tự động tại đây bằng Word)")
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

            add_main_heading(doc, "DANH MỤC CÁC SƠ ĐỒ")
            p_s = doc.add_paragraph("(Chèn danh mục sơ đồ tự động tại đây bằng Word)")
            p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # =====================================
            # SECTION 4: BẮT ĐẦU NỘI DUNG (SỐ Ả RẬP)
            # =====================================
            new_section_content = doc.add_section(WD_SECTION.NEW_PAGE)
            clear_page_border(new_section_content._sectPr)
            
            new_section_content.header.is_linked_to_previous = False
            header_para_main = new_section_content.header.paragraphs[0]
            header_para_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_page_number(header_para_main) 
            set_pgnum_type(new_section_content._sectPr, fmt='decimal', start='1')

            # Luôn bắt buộc xuất phần Đặt vấn đề
            add_main_heading(doc, "ĐẶT VẤN ĐỀ")
            if dat_van_de_content.strip():
                add_normal_text(doc, dat_van_de_content)
            doc.add_page_break()

            for i, chap in enumerate(all_chapters):
                add_main_heading(doc, f"Chương {i+1}. {chap['title']}")
                add_normal_text(doc, chap['content'])
                write_sections_to_word(doc, chap['children'], [str(i+1)])
                doc.add_page_break()

            if danh_muc_content.strip():
                add_main_heading(doc, "DANH MỤC CÁC CÔNG TRÌNH CÔNG BỐ CÓ LIÊN QUAN")
                add_normal_text(doc, danh_muc_content)
                doc.add_page_break()

            # =====================================
            # SECTION 5: TÀI LIỆU THAM KHẢO & PHỤ LỤC 
            # =====================================
            new_section_end = doc.add_section(WD_SECTION.CONTINUOUS)
            new_section_end.header.is_linked_to_previous = False
            for hp in new_section_end.header.paragraphs: hp.text = "" 

            add_main_heading(doc, "TÀI LIỆU THAM KHẢO")
            if tai_lieu_content.strip():
                add_normal_text(doc, tai_lieu_content)
            doc.add_page_break()

            add_main_heading(doc, "PHỤ LỤC")
            if phu_luc_content.strip():
                add_normal_text(doc, phu_luc_content)

            bio = io.BytesIO()
            doc.save(bio)
            
            st.success("🎉 Đã xuất file thành công!")
            st.info("💡 **LƯU Ý:** Thuật toán đã căn đều tự động khoảng cách giữa mốc ĐẠI HỌC và mốc THÀNH PHỐ. Các khối liên quan đã được dính liền nhau cực kỳ chuyên nghiệp!")
            st.download_button("⬇️ TẢI FILE ĐỀ CƯƠNG LUẬN VĂN (.docx)", bio.getvalue(), "De_Cuong_Hoan_Chinh_Can_Deu.docx", 
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
